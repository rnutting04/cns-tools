# app/services/letter_generator.py
import re
from io import BytesIO
from typing import Any

from docx import Document
from docx.oxml.ns import qn


def _replace_in_paragraph(paragraph, field_values: dict[str, Any]) -> None:
    full_text = "".join(run.text for run in paragraph.runs)
    if "{{" not in full_text:
        return

    for run in paragraph.runs:
        for key, value in field_values.items():
            placeholder = f"{{{{{key}}}}}"
            if placeholder in run.text:
                run.text = run.text.replace(
                    placeholder,
                    str(value) if value is not None else ""
                )


    remaining = "".join(run.text for run in paragraph.runs)
    if "{{" in remaining:
        raise ValueError(
            f"Unresolved placeholder (likely split across runs): {remaining}"
        )


def replace_placeholders(doc: Document, field_values: dict[str, Any]) -> Document:
    """Replace all {{placeholder}} tokens in paragraphs and table cells."""
    # Body paragraphs
    for para in doc.paragraphs:
        _replace_in_paragraph(para, field_values)

    # Table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para, field_values)

    # Headers and footers
    for section in doc.sections:
        for header_para in section.header.paragraphs:
            _replace_in_paragraph(header_para, field_values)
        for footer_para in section.footer.paragraphs:
            _replace_in_paragraph(footer_para, field_values)

    return doc


def generate_letter(template_bytes: bytes, field_values: dict[str, Any]) -> bytes:
    """Load a .docx from bytes, replace placeholders, return resulting bytes."""
    doc = Document(BytesIO(template_bytes))
    replace_placeholders(doc, field_values)
    output = BytesIO()
    doc.save(output)
    return output.getvalue()

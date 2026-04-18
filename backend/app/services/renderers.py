# app/services/renderers.py
from io import BytesIO

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from app.config.letter_config import (
    get_proxy_warning_text,
    render_proxy_vote,
)
from app.services.letter_generator import replace_placeholders


# ─── SIMPLE RENDERER ─────────────────────────────────────────

def simple_renderer(template_bytes: bytes, field_values: dict) -> bytes:
    doc = Document(BytesIO(template_bytes))
    replace_placeholders(doc, field_values)

    output = BytesIO()
    doc.save(output)
    return output.getvalue()


# ─── PROXY RENDERER ──────────────────────────────────────────

def _build_proxy_vote_table(doc: Document, vote_texts: list[str]):
    tmp = doc.add_table(rows=0, cols=2)

    for text in vote_texts:
        row = tmp.add_row().cells

        p_left = row[0].paragraphs[0]
        p_left.paragraph_format.space_before = Pt(0)
        p_left.paragraph_format.space_after = Pt(0)
        r_left = p_left.add_run(text)
        r_left.font.size = Pt(12)

        p_right = row[1].paragraphs[0]
        p_right.paragraph_format.space_before = Pt(0)
        p_right.paragraph_format.space_after = Pt(0)
        p_right.alignment = 1  # center
        r_right = p_right.add_run("YES ___ NO ___")
        r_right.font.size = Pt(12)

        _set_cell_width(row[0], 7200)
        _set_cell_width(row[1], 1800)

    _set_no_borders(tmp._tbl)

    tbl_el = tmp._tbl
    tbl_el.getparent().remove(tbl_el)
    return tbl_el


def proxy_renderer(template_bytes: bytes, field_values: dict) -> bytes:
    doc = Document(BytesIO(template_bytes))

    votes = field_values.get("votes", [])
    if not isinstance(votes, list):
        votes = []

    rendered_votes: list[str] = []
    needs_warning = False

    for vote in votes:
        if not isinstance(vote, dict):
            continue
        text, warning = render_proxy_vote(vote)
        if text:
            rendered_votes.append(text)
        if warning:
            needs_warning = True

    fv = dict(field_values)
    matter_count = len(rendered_votes)
    fv["proxy_matter_count"] = str(matter_count)
    fv["proxy_matter_count_word_num_lower"] = _word_num_lower(matter_count)

    # Remove block anchors before generic replacement
    vote_elements = []
    if rendered_votes:
        vote_elements.append(_build_proxy_vote_table(doc, rendered_votes))
    _replace_anchor(doc, "{{BLOCK:PROXY_VOTES}}", vote_elements)

    if needs_warning:
        warning_para = doc.add_paragraph()
        warning_run = warning_para.add_run(get_proxy_warning_text())
        warning_run.bold = True
        warning_run.font.size = Pt(12)
        _replace_anchor(doc, "{{BLOCK:PROXY_WARNING}}", [warning_para._element])
    else:
        _replace_anchor(doc, "{{BLOCK:PROXY_WARNING}}", [])

    replace_placeholders(doc, fv)

    output = BytesIO()
    doc.save(output)
    return output.getvalue()


# ─── COUNT WORD HELPERS ───────────────────────────────────────

_COUNT_WORDS = {
    1: "ONE", 2: "TWO", 3: "THREE", 4: "FOUR", 5: "FIVE",
    6: "SIX", 7: "SEVEN", 8: "EIGHT", 9: "NINE", 10: "TEN",
}


def _word_num(n: int) -> str:
    """Returns e.g. 'TWO(2)'"""
    word = _COUNT_WORDS.get(n, str(n))
    return f"{word}({n})"


def _word_num_lower(n: int) -> str:
    """Returns e.g. 'two (2)'"""
    word = _COUNT_WORDS.get(n, str(n)).lower()
    return f"{word} ({n})"


# ─── TABLE HELPERS ────────────────────────────────────────────

def _set_no_borders(tbl_element) -> None:
    """Strip all visible borders from a table element."""
    tblPr = tbl_element.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl_element.insert(0, tblPr)

    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)

    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "none")
        tblBorders.append(el)
    tblPr.append(tblBorders)


def _set_cell_width(cell, width_twips: int) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    existing = tcPr.find(qn("w:tcW"))
    if existing is not None:
        tcPr.remove(existing)
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(width_twips))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def _replace_anchor(doc: Document, anchor: str, new_elements: list) -> bool:
    """
    Find the first body paragraph containing `anchor`, replace it
    with `new_elements` (list of lxml elements). Returns True if found.
    """
    for para in doc.paragraphs:
        if anchor in para.text:
            parent = para._element.getparent()
            idx = list(parent).index(para._element)
            for i, el in enumerate(new_elements):
                parent.insert(idx + i, el)
            parent.remove(para._element)
            return True
    return False


# ─── BALLOT RENDERER ─────────────────────────────────────────
#
# Template placeholders:
#   {{candidate_count}}            – numeric count, e.g. "3"
#   {{candidate_count_word_num}}   – e.g. "THREE(3)"
#   {{candidate_count_word_num_lower}} – e.g. "three (3)"
#   {{BLOCK:BALLOT_CANDIDATES}}    – replaced with candidate table
#
# field_values["candidates"] must be a list of candidate full-name strings.


def _build_ballot_field_values(field_values: dict) -> tuple[list[str], dict]:
    candidates = list(field_values.get("candidates") or [])

    raw_limit = field_values.get("candidate_vote_limit", 0)
    try:
        vote_limit = int(raw_limit)
    except (TypeError, ValueError):
        vote_limit = 0

    fv = dict(field_values)
    fv["candidate_count"] = str(vote_limit)
    fv["candidate_count_word_num"] = _word_num(vote_limit)
    fv["candidate_count_word_num_lower"] = _word_num_lower(vote_limit)

    return candidates, fv


def _build_ballot_table(doc: Document, candidates: list[str], two_col_threshold: int = 8):
    """
    Build a centered, borderless ballot table.

    Behavior:
    - If candidate count <= two_col_threshold:
        one centered ballot column
    - If candidate count > two_col_threshold:
        two compact ballot columns to reduce height

    Each visible ballot entry is rendered like:
        ___   Jim Example
    """

    def _add_entry(cell, name: str) -> None:
        p = cell.paragraphs[0]
        p.alignment = 1  # center
        run = p.add_run(f"___   {name}")
        run.font.size = Pt(12)

    count = len(candidates)

    # One-column mode
    if count <= two_col_threshold:
        tmp = doc.add_table(rows=0, cols=1)

        for name in candidates:
            row = tmp.add_row().cells
            _add_entry(row[0], name)
            _set_cell_width(row[0], 5000)  # narrower, centered block

        _set_no_borders(tmp._tbl)

        # center the table itself
        tblPr = tmp._tbl.tblPr
        jc = tblPr.find(qn("w:jc"))
        if jc is None:
            jc = OxmlElement("w:jc")
            tblPr.append(jc)
        jc.set(qn("w:val"), "center")

        tbl_el = tmp._tbl
        tbl_el.getparent().remove(tbl_el)
        return tbl_el

    # Two-column mode
    left = candidates[: (count + 1) // 2]
    right = candidates[(count + 1) // 2 :]

    rows_needed = max(len(left), len(right))
    tmp = doc.add_table(rows=0, cols=2)

    for i in range(rows_needed):
        row = tmp.add_row().cells

        if i < len(left):
            _add_entry(row[0], left[i])
        else:
            row[0].text = ""

        if i < len(right):
            _add_entry(row[1], right[i])
        else:
            row[1].text = ""

        _set_cell_width(row[0], 3200)
        _set_cell_width(row[1], 3200)

    _set_no_borders(tmp._tbl)

    # center the whole table
    tblPr = tmp._tbl.tblPr
    jc = tblPr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        tblPr.append(jc)
    jc.set(qn("w:val"), "center")

    tbl_el = tmp._tbl
    tbl_el.getparent().remove(tbl_el)
    return tbl_el


def ballot_renderer(template_bytes: bytes, field_values: dict) -> bytes:

    candidates, fv = _build_ballot_field_values(field_values)

    doc = Document(BytesIO(template_bytes))

    if candidates:
        tbl_el = _build_ballot_table(doc, candidates)
        _replace_anchor(doc, "{{BLOCK:BALLOT_CANDIDATES}}", [tbl_el])
    else:
        _replace_anchor(doc, "{{BLOCK:BALLOT_CANDIDATES}}", [])

    replace_placeholders(doc, fv)

    output = BytesIO()
    doc.save(output)
    return output.getvalue()


# ─── ELECTRONIC BALLOT RENDERER ──────────────────────────────
#
# Template placeholders:
#   Same count placeholders as ballot renderer.
#   {{BLOCK:ELECTRONIC_BALLOT_CANDIDATES}} – replaced with candidate table
#
# field_values["candidates"] must be a list of candidate full-name strings.

def _build_electronic_ballot_table(doc: Document, candidates: list[str]):
    """
    Build a borderless 4-column table in a two-up ballot layout:

        1. Patricia Costa      ☐      5. Ed Malachowski      ☐
        2. Daniel Gilreath     ☐      6. Paul Melton         ☐

    Columns:
      1 = left numbered candidate
      2 = left checkbox
      3 = right numbered candidate
      4 = right checkbox

    Candidates are split evenly across the two sides.
    """
    tmp = doc.add_table(rows=0, cols=4)

    left = candidates[: (len(candidates) + 1) // 2]
    right = candidates[(len(candidates) + 1) // 2 :]

    rows_needed = max(len(left), len(right))

    for i in range(rows_needed):
        row = tmp.add_row().cells

        # Left side
        if i < len(left):
            p = row[0].paragraphs[0]
            p.alignment = 0  # left
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(f"{i + 1}. {left[i]}")
            r.font.size = Pt(12)

            p_box = row[1].paragraphs[0]
            p_box.alignment = 1  # center
            p_box.paragraph_format.space_before = Pt(0)
            p_box.paragraph_format.space_after = Pt(0)
            r_box = p_box.add_run("\u2610")  # ☐
            r_box.font.size = Pt(12)
        else:
            row[0].text = ""
            row[1].text = ""

        # Right side
        if i < len(right):
            right_num = len(left) + i + 1

            p = row[2].paragraphs[0]
            p.alignment = 0  # left
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(f"{right_num}. {right[i]}")
            r.font.size = Pt(12)

            p_box = row[3].paragraphs[0]
            p_box.alignment = 1  # center
            p_box.paragraph_format.space_before = Pt(0)
            p_box.paragraph_format.space_after = Pt(0)
            r_box = p_box.add_run("\u2610")  # ☐
            r_box.font.size = Pt(12)
        else:
            row[2].text = ""
            row[3].text = ""

        # Tight widths so the layout resembles the sample more closely
        _set_cell_width(row[0], 3000)  # left name
        _set_cell_width(row[1], 450)   # left checkbox
        _set_cell_width(row[2], 3000)  # right name
        _set_cell_width(row[3], 450)   # right checkbox

    _set_no_borders(tmp._tbl)

    # center the whole table
    tblPr = tmp._tbl.tblPr
    jc = tblPr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        tblPr.append(jc)
    jc.set(qn("w:val"), "center")

    tbl_el = tmp._tbl
    tbl_el.getparent().remove(tbl_el)
    return tbl_el


def electronic_ballot_renderer(template_bytes: bytes, field_values: dict) -> bytes:
    candidates, fv = _build_ballot_field_values(field_values)

    doc = Document(BytesIO(template_bytes))

    if candidates:
        tbl_el = _build_electronic_ballot_table(doc, candidates)
        _replace_anchor(doc, "{{BLOCK:ELECTRONIC_BALLOT_CANDIDATES}}", [tbl_el])
    else:
        _replace_anchor(doc, "{{BLOCK:ELECTRONIC_BALLOT_CANDIDATES}}", [])

    replace_placeholders(doc, fv)

    output = BytesIO()
    doc.save(output)
    return output.getvalue()


# ─── NOTICE / CANDIDACY RENDERER ─────────────────────────────
#
# Currently behaves like simple_renderer. Registered separately so
# the frontend and backend dispatch are both renderer-type-aware and
# ready for future customization without architectural changes.

def notice_candidacy_renderer(template_bytes: bytes, field_values: dict) -> bytes:
    doc = Document(BytesIO(template_bytes))
    replace_placeholders(doc, field_values)
    output = BytesIO()
    doc.save(output)
    return output.getvalue()


# ─── REGISTRY ────────────────────────────────────────────────

RENDERERS = {
    "simple": simple_renderer,
    "proxy": proxy_renderer,
    "ballot": ballot_renderer,
    "electronic_ballot": electronic_ballot_renderer,
    "notice_candidacy": notice_candidacy_renderer,
}
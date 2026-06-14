"""Helpers to build in-memory .docx files for renderer tests."""

from io import BytesIO

from docx import Document


def make_docx(paragraphs: list[str]) -> bytes:
    """
    Build a .docx whose body contains one paragraph per string, each placed in a
    single run so ``{{placeholder}}`` tokens are never split across runs.
    """
    doc = Document()
    for text in paragraphs:
        para = doc.add_paragraph()
        para.add_run(text)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def docx_text(docx_bytes: bytes) -> str:
    """Return all body-paragraph and table-cell text from a .docx as one string."""
    doc = Document(BytesIO(docx_bytes))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def docx_tables(docx_bytes: bytes):
    """Return the list of tables in a rendered .docx."""
    return Document(BytesIO(docx_bytes)).tables
